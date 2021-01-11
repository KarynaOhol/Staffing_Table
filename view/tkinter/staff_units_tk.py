import tkinter as tk
import tkinter.messagebox as mb
import tkinter.ttk as ttk
from controller.staff_units_controller import StaffUnitsController
from entity.staff_units import StaffUnits
from docx import Document
import openpyxl


class StaffUnitsApp(ttk.Frame):

    def __init__(self):
        super().__init__()
        self.stf_ctrl = StaffUnitsController()
        self.lblTitle = tk.Label(self, text="Staff Units",
                                 font=("Helvetica", 16), bg="yellow",
                                 fg="green")
        self.lblPosition = tk.Label(self, text="Enter Position:",
                                    font=("Helvetica", 10),
                                    bg="blue", fg="yellow")
        self.lblPersent = tk.Label(self, text="Enter persentage for harmful conditions:",
                                   font=("Helvetica", 10), bg="blue",
                                   fg="yellow")
        self.lblSalary = tk.Label(self, text="Enter salary:",
                                  font=("Helvetica", 10),
                                  bg="blue", fg="yellow")
        self.lblVacation = tk.Label(self, text="Enter vacation:",
                                    font=("Helvetica", 10),
                                    bg="blue", fg="yellow")
        self.lblSelect = tk.Label(self, text="Please select one record below to update or delete",
                                  font=("Helvetica", 10),
                                  bg="blue", fg="yellow")
        self.lblSearch = tk.Label(self, text="Please Enter Roll No:",
                                  font=("Helvetica", 10),
                                  bg="blue", fg="yellow")

        self.entPosition = tk.Entry(self)
        self.entPersent = tk.Entry(self)
        self.entSalary = tk.Entry(self)
        self.entVacation = tk.Entry(self)
        self.entSearch = tk.Entry(self)

        self.btn_xlsx = tk.Button(self, text="xlsx", font=("Helvetica", 11), bg="yellow", fg="blue",
                                  command=self.xlsx_export)
        self.btn_docx = tk.Button(self, text="docx", font=("Helvetica", 11), bg="yellow", fg="blue",
                                  command=self.docx_export)
        self.btn_register = tk.Button(self, text="Register", font=("Helvetica", 11), bg="yellow", fg="blue",
                                      command=self.register_staff_unit)
        self.btn_update = tk.Button(self, text="Update", font=("Helvetica", 11), bg="yellow", fg="blue",
                                    command=self.update_staff_unit)
        self.btn_delete = tk.Button(self, text="Delete", font=("Helvetica", 11), bg="yellow", fg="blue",
                                    command=self.delete_staff_unit)
        self.btn_clear = tk.Button(self, text="Clear", font=("Helvetica", 11), bg="yellow", fg="blue",
                                   command=self.clear_form)
        self.btn_show_all = tk.Button(self, text="Show All", font=("Helvetica", 11), bg="yellow", fg="blue",
                                      command=self.load_staff_units)
        self.btn_search = tk.Button(self, text="Search", font=("Helvetica", 11), bg="yellow", fg="blue",
                                    command=self.show_search_record)
        self.btn_exit = tk.Button(self, text="Exit", font=("Helvetica", 16), bg="yellow", fg="blue", command=self.exit)

        columns = ("#1", "#2", "#3", "#4", "#5")
        self.tvStaffUnit = ttk.Treeview(self, show="headings", height="5", columns=columns)
        self.tvStaffUnit.heading('#1', text='RollNo', anchor='center')
        self.tvStaffUnit.column('#1', width=60, anchor='center', stretch=False)
        self.tvStaffUnit.heading('#2', text='Position', anchor='center')
        self.tvStaffUnit.column('#2', width=10, anchor='center', stretch=True)
        self.tvStaffUnit.heading('#3', text='Percent', anchor='center')
        self.tvStaffUnit.column('#3', width=10, anchor='center', stretch=True)
        self.tvStaffUnit.heading('#4', text='Salary', anchor='center')
        self.tvStaffUnit.column('#4', width=10, anchor='center', stretch=False)
        self.tvStaffUnit.heading('#5', text='Vacation', anchor='center')
        self.tvStaffUnit.column('#5', width=10, anchor='center', stretch=False)

        vsb = ttk.Scrollbar(self, orient=tk.VERTICAL, command=self.tvStaffUnit.yview)
        vsb.place(x=40 + 640 + 1, y=310, height=180 + 20)
        self.tvStaffUnit.configure(yscroll=vsb.set)
        hsb = ttk.Scrollbar(self, orient=tk.HORIZONTAL, command=self.tvStaffUnit.xview)
        hsb.place(x=40, y=310 + 200 + 1, width=620 + 20)
        self.tvStaffUnit.configure(xscroll=hsb.set)
        self.tvStaffUnit.bind("<<TreeviewSelect>>", self.show_selected_record)

        self.lblTitle.place(x=280, y=30, height=27, width=300)
        self.lblPosition.place(x=175, y=70, height=23, width=100)
        self.lblPersent.place(x=20, y=100, height=23, width=250)
        self.lblSalary.place(x=171, y=129, height=23, width=104)
        self.lblVacation.place(x=170, y=158, height=23, width=100)
        self.lblSelect.place(x=150, y=280, height=23, width=400)
        self.lblSearch.place(x=174, y=560, height=23, width=134)
        self.entPosition.place(x=277, y=72, height=21, width=186)
        self.entPersent.place(x=277, y=100, height=21, width=186)
        self.entSalary.place(x=277, y=129, height=21, width=186)
        self.entVacation.place(x=277, y=158, height=21, width=186)
        self.entSearch.place(x=310, y=560, height=21, width=186)
        self.btn_xlsx.place(x=130, y=245, height=25, width=76)
        self.btn_docx.place(x=210, y=245, height=25, width=76)
        self.btn_register.place(x=290, y=245, height=25, width=76)
        self.btn_update.place(x=370, y=245, height=25, width=76)
        self.btn_delete.place(x=460, y=245, height=25, width=76)
        self.btn_clear.place(x=548, y=245, height=25, width=76)
        self.btn_show_all.place(x=630, y=245, height=25, width=76)
        self.btn_search.place(x=498, y=558, height=26, width=60)
        self.btn_exit.place(x=320, y=610, height=31, width=60)
        self.tvStaffUnit.place(x=40, y=310, height=200, width=640)
        self.load_staff_units()

    def clear_form(self):
        self.entPosition.delete(0, tk.END)
        self.entPersent.delete(0, tk.END)
        self.entSalary.delete(0, tk.END)
        self.entVacation.delete(0, tk.END)

    def exit(self):
        MsgBox = mb.askquestion('Exit Application',
                                'Are you sure you want to exit the application',
                                icon='warning')
        if MsgBox == 'yes':
            self.destroy()

    def delete_staff_unit(self):
        MsgBox = mb.askquestion('Delete Record', 'Are you sure! you want to delete selected record',
                                icon='warning')

        if MsgBox == 'yes':
            self.stf_ctrl.delete(roll_no)
            mb.showinfo("Information", "Record Deleted Succssfully")
            self.load_staff_units()
            self.clear_form()

    def register_staff_unit(self):

        position = self.entPosition.get()
        pers = self.entPersent.get()
        sal = self.entSalary.get()
        vac = self.entVacation.get()

        # validating Entry Widgets
        if position == "":
            mb.showinfo('Information', "Please Enter position")
            self.entPosition.focus_set()
            return
        if pers == "":
            mb.showinfo('Information', "Please Enter persent")
            self.entPersent.focus_set()
            return
        if sal == "":
            mb.showinfo('Information', "Please Enter salary")
            self.entSalary.focus_set()
            return
        if vac == "":
            mb.showinfo('Information', "Please Enter vacation")
            self.entVacation.focus_set()
            return
        # Inserting record
        try:
            self.stf_ctrl.create(StaffUnits(pers, vac, position, sal))
            self.load_staff_units()
        except Exception as err:
            print(err)

    def show_search_record(self):
        s_roll_no = self.entSearch.get()
        print(s_roll_no)
        if s_roll_no == "":
            mb.showinfo('Information', "Please Enter Roll")
            self.entSearch.focus_set()
            return

        self.tvStaffUnit.delete(*self.tvStaffUnit.get_children())
        stf = self.stf_ctrl.get_by_id(s_roll_no)
        self.tvStaffUnit.insert("", 'end', text="StaffUnits",
                                values=(stf.id, stf.position, stf.persentage_two, stf.salary, stf.vacation))

    def show_selected_record(self, event):
        self.clear_form()
        for selection in self.tvStaffUnit.selection():
            item = self.tvStaffUnit.item(selection)
        global roll_no
        roll_no, pos, pers, sal, vac = item["values"][0:5]
        self.entPosition.insert(0, pos)
        self.entPersent.insert(0, pers)
        self.entSalary.insert(0, sal)
        self.entVacation.insert(0, vac)
        return roll_no

    def update_staff_unit(self):
        position = self.entPosition.get()
        pers = self.entPersent.get()
        sal = self.entSalary.get()
        vac = self.entVacation.get()
        print(roll_no)
        self.stf_ctrl.update(roll_no, StaffUnits(pers, vac, position, sal))
        mb.showinfo("Info", "Selected Record Updated Successfully ")
        self.load_staff_units()

    def load_staff_units(self):

        self.tvStaffUnit.delete(*self.tvStaffUnit.get_children())
        units = self.stf_ctrl.all_units()
        for u in units:
            self.tvStaffUnit.insert("", 'end', text="StaffUnits",
                                    values=(u.id, u.position, u.persentage_two, u.salary, u.vacation))

    def xlsx_export(self):
        wb = openpyxl.load_workbook('C:\\Users\\karina\\PycharmProjects\\staffing\\reports\\All.xlsx')
        if 'Staff units' not in wb.sheetnames:
            wb.create_sheet('Staff units')
        ws = wb.get_sheet_by_name('Staff units')
        ws.delete_cols(1, 5)
        ws.delete_rows(1, 100)
        i = 1
        ws.cell(row=i, column=1).value = "ID"
        ws.cell(row=i, column=2).value = "Persent2"
        ws.cell(row=i, column=3).value = "Vacation"
        ws.cell(row=i, column=4).value = "Position"
        ws.cell(row=i, column=5).value = "Salary"
        units = self.stf_ctrl.all_units()
        for u in units:
            ws.cell(row=i + 1, column=1).value = u.id
            ws.cell(row=i + 1, column=2).value = u.persentage_two
            ws.cell(row=i + 1, column=3).value = u.vacation
            ws.cell(row=i + 1, column=4).value = u.position
            ws.cell(row=i + 1, column=5).value = u.salary
            i += 1
        wb.save('C:\\Users\\karina\\PycharmProjects\\staffing\\reports\\All.xlsx')
        print("xlsx successful")

    def docx_export(self):
        document = Document()
        document.add_heading("Staff units", 0)
        table = document.add_table(rows=1, cols=5)
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'ID'
        hdr_cells[1].text = 'Persent for harmful conditions'
        hdr_cells[2].text = 'Vacation'
        hdr_cells[3].text = 'Position'
        hdr_cells[4].text = 'Salary'
        units = self.stf_ctrl.all_units()
        for u in units:
            row_cells = table.add_row().cells
            row_cells[0].text = str(u.id)
            row_cells[1].text = str(u.persentage_two)
            row_cells[2].text = str(u.vacation)
            row_cells[3].text = u.position
            row_cells[4].text = str(u.salary)
        document.save('C:\\Users\\karina\\PycharmProjects\\staffing\\reports\\Staff Units.docx')
        print("docx successful")

import tkinter as tk
import tkinter.messagebox as mb
import tkinter.ttk as ttk
from controller.divisions_controller import DivisionsController
from entity.divisions import Division
from docx import Document
import openpyxl


class DivisionApp(ttk.Frame):

    def __init__(self):
        super().__init__()
        self.div_ctrl = DivisionsController()
        self.lblTitle = tk.Label(self, text="Divisions",
                                 font=("Helvetica", 16), bg="yellow",
                                 fg="green")
        self.lblName = tk.Label(self, text="Enter Name:",
                                font=("Helvetica", 10),
                                bg="blue", fg="yellow")
        self.lblPersent = tk.Label(self, text="Enter persentage per irregular day:",
                                   font=("Helvetica", 10), bg="blue",
                                   fg="yellow")
        self.lblType = tk.Label(self, text="Enter type:",
                                font=("Helvetica", 10),
                                bg="blue", fg="yellow")
        self.lblSelect = tk.Label(self, text="Please select one record below to update or delete",
                                  font=("Helvetica", 10),
                                  bg="blue", fg="yellow")
        self.lblSearch = tk.Label(self, text="Please Enter Roll No:",
                                  font=("Helvetica", 10),
                                  bg="blue", fg="yellow")

        self.entName = tk.Entry(self)
        self.entPersent = tk.Entry(self)
        self.entType = tk.Entry(self)
        self.entSearch = tk.Entry(self)

        self.btn_xlsx = tk.Button(self, text="xlsx", font=("Helvetica", 11), bg="yellow", fg="blue",
                                  command=self.xlsx_export)
        self.btn_docx = tk.Button(self, text="docx", font=("Helvetica", 11), bg="yellow", fg="blue",
                                  command=self.docx_export)
        self.btn_register = tk.Button(self, text="Register", font=("Helvetica", 11), bg="yellow", fg="blue",
                                      command=self.register_division)
        self.btn_update = tk.Button(self, text="Update", font=("Helvetica", 11), bg="yellow", fg="blue",
                                    command=self.update_division)
        self.btn_delete = tk.Button(self, text="Delete", font=("Helvetica", 11), bg="yellow", fg="blue",
                                    command=self.delete_division)
        self.btn_clear = tk.Button(self, text="Clear", font=("Helvetica", 11), bg="yellow", fg="blue",
                                   command=self.clear_form)
        self.btn_show_all = tk.Button(self, text="Show All", font=("Helvetica", 11), bg="yellow", fg="blue",
                                      command=self.load_divisions)
        self.btn_search = tk.Button(self, text="Search", font=("Helvetica", 11), bg="yellow", fg="blue",
                                    command=self.show_search_record)
        self.btn_exit = tk.Button(self, text="Exit", font=("Helvetica", 16), bg="yellow", fg="blue", command=self.exit)

        columns = ("#1", "#2", "#3", "#4")
        self.tvDivision = ttk.Treeview(self, show="headings", height="5", columns=columns)
        self.tvDivision.heading('#1', text='RollNo', anchor='center')
        self.tvDivision.column('#1', width=60, anchor='center', stretch=False)
        self.tvDivision.heading('#2', text='Persentage', anchor='center')
        self.tvDivision.column('#2', width=10, anchor='center', stretch=True)
        self.tvDivision.heading('#3', text='Type', anchor='center')
        self.tvDivision.column('#3', width=10, anchor='center', stretch=True)
        self.tvDivision.heading('#4', text='Name', anchor='center')
        self.tvDivision.column('#4', width=10, anchor='center', stretch=False)

        vsb = ttk.Scrollbar(self, orient=tk.VERTICAL, command=self.tvDivision.yview)
        vsb.place(x=40 + 640 + 1, y=310, height=180 + 20)
        self.tvDivision.configure(yscroll=vsb.set)
        hsb = ttk.Scrollbar(self, orient=tk.HORIZONTAL, command=self.tvDivision.xview)
        hsb.place(x=40, y=310 + 200 + 1, width=620 + 20)
        self.tvDivision.configure(xscroll=hsb.set)
        self.tvDivision.bind("<<TreeviewSelect>>", self.show_selected_record)

        self.lblTitle.place(x=280, y=30, height=27, width=300)
        self.lblName.place(x=175, y=70, height=23, width=100)
        self.lblPersent.place(x=20, y=100, height=23, width=250)
        self.lblType.place(x=171, y=129, height=23, width=104)
        self.lblSelect.place(x=150, y=280, height=23, width=400)
        self.lblSearch.place(x=174, y=560, height=23, width=134)
        self.entName.place(x=277, y=72, height=21, width=186)
        self.entPersent.place(x=277, y=100, height=21, width=186)
        self.entType.place(x=277, y=129, height=21, width=186)
        self.entSearch.place(x=310, y=560, height=21, width=186)
        self.btn_xlsx.place(x=130, y=245, height=25, width=76)
        self.btn_docx.place(x=210, y=245, height=25, width=76)
        self.btn_register.place(x=290, y=245, height=25, width=76)
        self.btn_update.place(x=370, y=245, height=25, width=76)
        self.btn_delete.place(x=460, y=245, height=25, width=76)
        self.btn_clear.place(x=548, y=245, height=25, width=76)
        self.btn_show_all.place(x=630, y=245, height=25, width=76)
        self.btn_search.place(x=498, y=558, height=26, width=60)
        self.btn_exit.place(x=320, y=610, height=31, width=60)
        self.tvDivision.place(x=40, y=310, height=200, width=640)
        self.load_divisions()

    def clear_form(self):
        self.entName.delete(0, tk.END)
        self.entPersent.delete(0, tk.END)
        self.entType.delete(0, tk.END)

    def exit(self):
        MsgBox = mb.askquestion('Exit Application',
                                'Are you sure you want to exit the application',
                                icon='warning')
        if MsgBox == 'yes':
            self.destroy()

    def delete_division(self):
        MsgBox = mb.askquestion('Delete Record', 'Are you sure! you want to delete selected record',
                                icon='warning')

        if MsgBox == 'yes':
            self.div_ctrl.delete(roll_no)
            mb.showinfo("Information", "Record Deleted Succssfully")
            self.load_divisions()
            self.clear_form()

    def register_division(self):

        name = self.entName.get()  # Retrieving entered first name
        pers = self.entPersent.get()  # Retrieving entered last name
        type = self.entType.get()  # Retrieving entered contact number

        # validating Entry Widgets
        if name == "":
            mb.showinfo('Information', "Please Enter name")
            self.entName.focus_set()
            return
        if pers == "":
            mb.showinfo('Information', "Please Enter persent")
            self.entPersent.focus_set()
            return
        if type == "":
            mb.showinfo('Information', "Please Enter type")
            self.entType.focus_set()
            return
        # Inserting record
        try:
            self.div_ctrl.create(Division(name, pers, type))
            self.load_divisions()
        except Exception as err:
            print(err)

    def show_search_record(self):
        s_roll_no = self.entSearch.get()  # Retrieving entered first name
        print(s_roll_no)
        if s_roll_no == "":
            mb.showinfo('Information', "Please Enter Roll")
            self.entSearch.focus_set()
            return

        self.tvDivision.delete(*self.tvDivision.get_children())
        division = self.div_ctrl.get_by_id(s_roll_no)
        self.tvDivision.insert("", 'end', text="Division",
                               values=(division.id, division.name, division.persentage_one, division.type))

    def show_selected_record(self, event):
        self.clear_form()
        for selection in self.tvDivision.selection():
            item = self.tvDivision.item(selection)
        global roll_no
        roll_no, name, pers, type = item["values"][0:4]
        self.entName.insert(0, name)
        self.entPersent.insert(0, pers)
        self.entType.insert(0, type)
        return roll_no

    def update_division(self):
        name = self.entName.get()
        pers = self.entPersent.get()
        type = self.entType.get()
        print(roll_no)
        self.div_ctrl.update(roll_no, Division(name, pers, type))
        mb.showinfo("Info", "Selected  Record Updated Successfully ")
        self.load_divisions()

    def load_divisions(self):
        self.tvDivision.delete(*self.tvDivision.get_children())
        divisions = self.div_ctrl.all_divisions()
        for div in divisions:
            self.tvDivision.insert("", 'end', text="Divisions",
                                   values=(div.id, div.persentage_one, div.type, div.name))

    def xlsx_export(self):
        wb = openpyxl.load_workbook('C:\\Users\\karina\\PycharmProjects\\staffing\\reports\\All.xlsx')
        if 'Divisions' not in wb.sheetnames:
            wb.create_sheet('Divisions')
        ws = wb.get_sheet_by_name('Divisions')
        ws.delete_cols(1, 4)
        ws.delete_rows(1, 100)
        i = 1
        ws.cell(row=i, column=1).value = "ID"
        ws.cell(row=i, column=2).value = "Persent1"
        ws.cell(row=i, column=3).value = "Type"
        ws.cell(row=i, column=4).value = "Name"
        divisions = self.div_ctrl.all_divisions()
        for div in divisions:
            ws.cell(row=i + 1, column=1).value = div.id
            ws.cell(row=i + 1, column=2).value = div.persentage_one
            ws.cell(row=i + 1, column=3).value = div.type
            ws.cell(row=i + 1, column=4).value = div.name
            i += 1
        wb.save('C:\\Users\\karina\\PycharmProjects\\staffing\\reports\\All.xlsx')
        print("xlsx successful")

    def docx_export(self):
        document = Document()
        document.add_heading("Divisions", 0)
        table = document.add_table(rows=1, cols=4)
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'ID'
        hdr_cells[1].text = 'Name'
        hdr_cells[2].text = 'Persentage per irregular day'
        hdr_cells[3].text = 'Type'
        divisions = self.div_ctrl.all_divisions()
        for div in divisions:
            row_cells = table.add_row().cells
            row_cells[0].text = str(div.id)
            row_cells[1].text = div.name
            row_cells[2].text = str(div.persentage_one)
            row_cells[3].text = div.type
        document.save('C:\\Users\\karina\\PycharmProjects\\staffing\\reports\\Divisions.docx')
        print("docx successful")

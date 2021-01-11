import tkinter as tk
import tkinter.messagebox as mb
import tkinter.ttk as ttk
from controller.distributions_controller import DistributionsController
from entity.distribution import Distribution
from docx import Document
import openpyxl


class DistributionApp(ttk.Frame):

    def __init__(self):
        super().__init__()
        self.dst_ctrl = DistributionsController()
        self.lblTitle = tk.Label(self, text="Distribution",
                                 font=("Helvetica", 16), bg="yellow",
                                 fg="green")
        self.lblDivisionId = tk.Label(self, text="Enter division id:",
                                      font=("Helvetica", 10),
                                      bg="blue", fg="yellow")
        self.lblUnitId = tk.Label(self, text="Enter staff unit id:",
                                  font=("Helvetica", 10), bg="blue",
                                  fg="yellow")
        self.lblSelect = tk.Label(self, text="Please select one record below to update or delete",
                                  font=("Helvetica", 10),
                                  bg="blue", fg="yellow")
        self.lblSearch = tk.Label(self, text="Please Enter Roll No:",
                                  font=("Helvetica", 10),
                                  bg="blue", fg="yellow")
        self.lblSalary = tk.Label(self, text="Salary",
                                  font=("Helvetica", 10),
                                  bg="blue", fg="yellow")

        self.entDivisionId = tk.Entry(self)
        self.entUnitId = tk.Entry(self)
        self.entSearch = tk.Entry(self)

        self.btn_xlsx = tk.Button(self, text="xlsx", font=("Helvetica", 11), bg="yellow", fg="blue",
                                  command=self.xlsx_export)
        self.btn_docx = tk.Button(self, text="docx", font=("Helvetica", 11), bg="yellow", fg="blue",
                                  command=self.docx_export)
        self.btn_calc = tk.Button(self, text="Salary", font=("Helvetica", 11), bg="yellow", fg="blue",
                                  command=self.calc)
        self.btn_register = tk.Button(self, text="Register", font=("Helvetica", 11), bg="yellow", fg="blue",
                                      command=self.register_distribution)
        self.btn_update = tk.Button(self, text="Update", font=("Helvetica", 11), bg="yellow", fg="blue",
                                    command=self.update_staff_unit)
        self.btn_delete = tk.Button(self, text="Delete", font=("Helvetica", 11), bg="yellow", fg="blue",
                                    command=self.delete_distribution)
        self.btn_clear = tk.Button(self, text="Clear", font=("Helvetica", 11), bg="yellow", fg="blue",
                                   command=self.clear_form)
        self.btn_show_all = tk.Button(self, text="Show All", font=("Helvetica", 11), bg="yellow", fg="blue",
                                      command=self.load_distributions)
        self.btn_search = tk.Button(self, text="Search", font=("Helvetica", 11), bg="yellow", fg="blue",
                                    command=self.show_search_record)
        self.btn_exit = tk.Button(self, text="Exit", font=("Helvetica", 16), bg="yellow", fg="blue", command=self.exit)

        columns = ("#1", "#2", "#3", "#4", "#5", "#6", "#7")
        self.tvDistribution = ttk.Treeview(self, show="headings", height="5", columns=columns)
        self.tvDistribution.heading('#1', text='Staff Unit Id', anchor='center')
        self.tvDistribution.column('#1', width=70, anchor='center', stretch=False)
        self.tvDistribution.heading('#2', text='Position', anchor='center')
        self.tvDistribution.column('#2', width=1, anchor='center', stretch=True)
        self.tvDistribution.heading('#3', text='Percent2', anchor='center')
        self.tvDistribution.column('#3', width=1, anchor='center', stretch=True)
        self.tvDistribution.heading('#4', text='Salary', anchor='center')
        self.tvDistribution.column('#4', width=50, anchor='center', stretch=False)
        self.tvDistribution.heading('#5', text='Division Id', anchor='center')
        self.tvDistribution.column('#5', width=70, anchor='center', stretch=False)
        self.tvDistribution.heading('#6', text='Name', anchor='center')
        self.tvDistribution.column('#6', width=20, anchor='center', stretch=True)
        self.tvDistribution.heading('#7', text='Percent1', anchor='center')
        self.tvDistribution.column('#7', width=10, anchor='center', stretch=False)

        vsb = ttk.Scrollbar(self, orient=tk.VERTICAL, command=self.tvDistribution.yview)
        vsb.place(x=40 + 640 + 1, y=310, height=180 + 20)
        self.tvDistribution.configure(yscroll=vsb.set)
        hsb = ttk.Scrollbar(self, orient=tk.HORIZONTAL, command=self.tvDistribution.xview)
        hsb.place(x=40, y=310 + 200 + 1, width=620 + 20)
        self.tvDistribution.configure(xscroll=hsb.set)
        self.tvDistribution.bind("<<TreeviewSelect>>", self.show_selected_record)

        self.lblTitle.place(x=280, y=30, height=27, width=300)
        self.lblDivisionId.place(x=175, y=70, height=23, width=100)
        self.lblUnitId.place(x=20, y=100, height=23, width=250)
        self.lblSelect.place(x=150, y=280, height=23, width=400)
        self.lblSalary.place(x=174, y=530, height=23, width=300)
        self.lblSearch.place(x=174, y=560, height=23, width=134)
        self.entDivisionId.place(x=277, y=72, height=21, width=186)
        self.entUnitId.place(x=277, y=100, height=21, width=186)
        self.entSearch.place(x=310, y=560, height=21, width=186)
        self.btn_xlsx.place(x=50, y=245, height=25, width=76)
        self.btn_docx.place(x=130, y=245, height=25, width=76)
        self.btn_calc.place(x=210, y=245, height=25, width=76)
        self.btn_register.place(x=290, y=245, height=25, width=76)
        self.btn_update.place(x=370, y=245, height=25, width=76)
        self.btn_delete.place(x=460, y=245, height=25, width=76)
        self.btn_clear.place(x=548, y=245, height=25, width=76)
        self.btn_show_all.place(x=630, y=245, height=25, width=76)
        self.btn_search.place(x=498, y=558, height=26, width=60)
        self.btn_exit.place(x=320, y=610, height=31, width=60)
        self.tvDistribution.place(x=40, y=310, height=200, width=640)
        self.load_distributions()

    def clear_form(self):
        self.entDivisionId.delete(0, tk.END)
        self.entUnitId.delete(0, tk.END)

    def exit(self):
        MsgBox = mb.askquestion('Exit Application',
                                'Are you sure you want to exit the application',
                                icon='warning')
        if MsgBox == 'yes':
            self.destroy()

    def delete_distribution(self):
        MsgBox = mb.askquestion('Delete Record', 'Are you sure! you want to delete selected record',
                                icon='warning')

        if MsgBox == 'yes':
            self.dst_ctrl.delete(roll_no)
            mb.showinfo("Information", "Record Deleted Succssfully")
            self.load_distributions()
            self.clear_form()

    def register_distribution(self):
        division_id = self.entDivisionId.get()
        unit_id = self.entUnitId.get()

        # validating Entry Widgets
        if division_id == "":
            mb.showinfo('Information', "Please Enter division_id")
            self.entDivisionId.focus_set()
            return
        if unit_id == "":
            mb.showinfo('Information', "Please Enter persent")
            self.entUnitId.focus_set()
            return
        # Inserting record
        try:
            self.dst_ctrl.create(unit_id, division_id)
            self.load_distributions()
        except Exception as err:
            print(err)

    def show_search_record(self):
        s_roll_no = self.entSearch.get()
        print(s_roll_no)
        if s_roll_no == "":
            mb.showinfo('Information', "Please Enter unit id")
            self.entSearch.focus_set()
            return

        self.tvDistribution.delete(*self.tvDistribution.get_children())
        division = self.dst_ctrl.get_division_by_unit_id(s_roll_no)
        self.tvDistribution.insert("", 'end', text="StaffUnits",
                                   values=(s_roll_no, '', '', '', division.id, division.name,
                                           division.persentage_one))

    def show_selected_record(self, event):
        self.clear_form()
        for selection in self.tvDistribution.selection():
            item = self.tvDistribution.item(selection)
        global roll_no
        roll_no, pos, pers2, sal, div, nam, pers1 = item["values"][0:7]
        self.entDivisionId.insert(0, div)
        self.entUnitId.insert(0, roll_no)
        return roll_no

    def update_staff_unit(self):
        div_id = self.entDivisionId.get()
        print(roll_no)
        self.dst_ctrl.update(roll_no, div_id)
        mb.showinfo("Info", "Selected Record Updated Successfully ")
        self.load_distributions()

    def calc(self):
        if roll_no == "":
            mb.showinfo('Information', "Please select record")
            return
        res = self.dst_ctrl.calculate_salary_for_unit(roll_no)
        self.lblSalary.config(text=res)

    def load_distributions(self):

        self.tvDistribution.delete(*self.tvDistribution.get_children())
        dist = self.dst_ctrl.all_distributions()
        for d in dist:
            self.tvDistribution.insert("", 'end', text="StaffUnits",
                                       values=(d.staff_unit.id, d.staff_unit.position,
                                               d.staff_unit.persentage_two,
                                               d.staff_unit.salary, d.division.id,
                                               d.division.name, d.division.persentage_one))
    def xlsx_export(self):
        wb = openpyxl.load_workbook('C:\\Users\\karina\\PycharmProjects\\staffing\\reports\\All.xlsx')
        if 'Distribution units' not in wb.sheetnames:
            wb.create_sheet('Distribution units')
        ws = wb.get_sheet_by_name('Distribution units')
        ws.delete_cols(1, 7)
        ws.delete_rows(1, 100)
        i = 1
        ws.cell(row=i, column=1).value = "Division Id"
        ws.cell(row=i, column=2).value = "Name"
        ws.cell(row=i, column=3).value = "Persent 1"
        ws.cell(row=i, column=4).value = "Unit Id"
        ws.cell(row=i, column=5).value = "Persent 2"
        ws.cell(row=i, column=6).value = "Position"
        ws.cell(row=i, column=7).value = "Salary"
        dst = self.dst_ctrl.all_distributions()
        for u in dst:
            ws.cell(row=i + 1, column=1).value = u.division.id
            ws.cell(row=i + 1, column=2).value = u.division.name
            ws.cell(row=i + 1, column=3).value = u.division.persentage_one
            ws.cell(row=i + 1, column=4).value = u.staff_unit.id
            ws.cell(row=i + 1, column=5).value = u.staff_unit.persentage_two
            ws.cell(row=i + 1, column=6).value = u.staff_unit.position
            ws.cell(row=i + 1, column=7).value = u.staff_unit.salary
            i += 1
        wb.save('C:\\Users\\karina\\PycharmProjects\\staffing\\reports\\All.xlsx')
        print("xlsx successful")

    def docx_export(self):
        document = Document()
        document.add_heading("Distribution units", 0)
        table = document.add_table(rows=1, cols=7)
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Division Id'
        hdr_cells[1].text = 'Name'
        hdr_cells[2].text = 'Persentage per irregular day'
        hdr_cells[3].text = 'Unit Id'
        hdr_cells[4].text = 'Persent for harmful conditions'
        hdr_cells[5].text = 'Position'
        hdr_cells[6].text = 'Salary'
        dst = self.dst_ctrl.all_distributions()
        for u in dst:
            row_cells = table.add_row().cells
            row_cells[0].text = str(u.division.id)
            row_cells[1].text = u.division.name
            row_cells[2].text = str(u.division.persentage_one)
            row_cells[3].text = str(u.staff_unit.id)
            row_cells[4].text = str(u.staff_unit.persentage_two)
            row_cells[5].text = u.staff_unit.position
            row_cells[6].text = str(u.staff_unit.salary)
        document.save('C:\\Users\\karina\\PycharmProjects\\staffing\\reports\\Distributions.docx')
        print("docx successful")

import sys
from PyQt5 import QtWidgets, QtGui
from view.PyOt import design
from controller.distributions_controller import DistributionsController
from entity.distribution import Distribution
from controller.divisions_controller import DivisionsController
from entity.divisions import Division
from controller.staff_units_controller import StaffUnitsController
from entity.staff_units import StaffUnits

from docx import Document
import openpyxl


def get_selected_record(table):
    row = table.currentRow()
    roll_no = table.item(row, 0).text()
    return roll_no


class MainWindow(QtWidgets.QMainWindow, design.Ui_MainWindow):
    def __init__(self):
        super().__init__()
        self.setupUi(self)
        self.dst_ctrl = DistributionsController()
        self.div_ctrl = DivisionsController()
        self.stf_ctrl = StaffUnitsController()

        ####============DST
        self.btnXlsx.clicked.connect(self.xlsx_export_dst)
        self.btnDocx.clicked.connect(self.docx_export_dst)
        self.btnSalary.clicked.connect(self.calc)
        self.btnReg.clicked.connect(self.register_distribution)
        self.bntUpdate.clicked.connect(self.update_staff_unit_dst)
        self.btnDel.clicked.connect(self.delete_distribution)
        self.btnClear.clicked.connect(self.clear_form_dst)
        self.btnAll.clicked.connect(self.load_distributions)
        self.btnSearch.clicked.connect(self.show_search_record_dst)
        self.tableWidget.itemSelectionChanged.connect(self.set_ent_dst)
        self.load_distributions()
        ####============DIV
        self.btnXlsx_3.clicked.connect(self.xlsx_export_div)
        self.btnDocx_3.clicked.connect(self.docx_export_div)
        self.btnReg_3.clicked.connect(self.register_division)
        self.bntUpdate_3.clicked.connect(self.update_division)
        self.btnDel_3.clicked.connect(self.delete_division)
        self.btnClear_3.clicked.connect(self.clear_form_div)
        self.btnAll_3.clicked.connect(self.load_divisions)
        self.btnSearch_3.clicked.connect(self.show_search_record_div)
        self.tableWidget_3.itemSelectionChanged.connect(self.set_ent_div)
        self.load_divisions()
        ####============STF
        self.btnXlsx_8.clicked.connect(self.xlsx_export_stf)
        self.btnDocx_8.clicked.connect(self.docx_export_stf)
        self.btnReg_8.clicked.connect(self.register_staff_unit)
        self.bntUpdate_8.clicked.connect(self.update_staff_unit)
        self.btnDel_8.clicked.connect(self.delete_staff_unit)
        self.btnClear_8.clicked.connect(self.clear_form_stf)
        self.btnAll_8.clicked.connect(self.load_staff_units)
        self.btnSearch_8.clicked.connect(self.show_search_record_stf)
        self.tableWidget_8.itemSelectionChanged.connect(self.set_ent_stf)
        self.load_staff_units()

    ####================================================================DST
    def xlsx_export_dst(self):
        wb = openpyxl.load_workbook('C:\\Users\\karina\\PycharmProjects\\staffing\\reports\\All.xlsx')
        if 'Distribution units' not in wb.sheetnames:
            wb.create_sheet('Distribution units')
        ws = wb.get_sheet_by_name('Distribution units')
        ws.delete_cols(1, 7)
        ws.delete_rows(1, 100)
        i = 1
        ws.cell(row=i, column=1).value = "Division Id"
        ws.cell(row=i, column=2).value = "Name"
        ws.cell(row=i, column=3).value = "Persent 1"
        ws.cell(row=i, column=4).value = "Unit Id"
        ws.cell(row=i, column=5).value = "Persent 2"
        ws.cell(row=i, column=6).value = "Position"
        ws.cell(row=i, column=7).value = "Salary"
        dst = self.dst_ctrl.all_distributions()
        for u in dst:
            ws.cell(row=i + 1, column=1).value = u.division.id
            ws.cell(row=i + 1, column=2).value = u.division.name
            ws.cell(row=i + 1, column=3).value = u.division.persentage_one
            ws.cell(row=i + 1, column=4).value = u.staff_unit.id
            ws.cell(row=i + 1, column=5).value = u.staff_unit.persentage_two
            ws.cell(row=i + 1, column=6).value = u.staff_unit.position
            ws.cell(row=i + 1, column=7).value = u.staff_unit.salary
            i += 1
        wb.save('C:\\Users\\karina\\PycharmProjects\\staffing\\reports\\All.xlsx')
        print("xlsx successful")

    def docx_export_dst(self):
        document = Document()
        document.add_heading("Distribution units", 0)
        table = document.add_table(rows=1, cols=7)
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Division Id'
        hdr_cells[1].text = 'Name'
        hdr_cells[2].text = 'Persentage per irregular day'
        hdr_cells[3].text = 'Unit Id'
        hdr_cells[4].text = 'Persent for harmful conditions'
        hdr_cells[5].text = 'Position'
        hdr_cells[6].text = 'Salary'
        dst = self.dst_ctrl.all_distributions()
        for u in dst:
            row_cells = table.add_row().cells
            row_cells[0].text = str(u.division.id)
            row_cells[1].text = u.division.name
            row_cells[2].text = str(u.division.persentage_one)
            row_cells[3].text = str(u.staff_unit.id)
            row_cells[4].text = str(u.staff_unit.persentage_two)
            row_cells[5].text = u.staff_unit.position
            row_cells[6].text = str(u.staff_unit.salary)
        document.save('C:\\Users\\karina\\PycharmProjects\\staffing\\reports\\Distributions.docx')
        print("docx successful")

    def register_distribution(self):
        division_id = self.entDivID.text()
        unit_id = self.entStaffID.text()
        try:
            self.dst_ctrl.create(unit_id, division_id)
            self.load_distributions()
        except Exception as err:
            print(err)

    def update_staff_unit_dst(self):
        div_id = self.entDivID.text()
        print(get_selected_record(self.tableWidget))
        self.dst_ctrl.update(get_selected_record(self.tableWidget), div_id)
        self.load_distributions()

    def delete_distribution(self):
        self.dst_ctrl.delete(get_selected_record(self.tableWidget))
        self.load_distributions()
        self.clear_form_dst()

    def clear_form_dst(self):
        self.entDivID.clear()
        self.entStaffID.clear()
        self.entID.clear()

    def load_distributions(self):
        self.tableWidget.clearContents()
        dist = self.dst_ctrl.all_distributions()
        i = 0
        self.tableWidget.setRowCount(dist.__len__())
        for d in dist:
            self.tableWidget.setItem(i, 0, QtWidgets.QTableWidgetItem(str(d.staff_unit.id)))
            self.tableWidget.setItem(i, 1, QtWidgets.QTableWidgetItem(d.staff_unit.position))
            self.tableWidget.setItem(i, 2, QtWidgets.QTableWidgetItem(str(d.staff_unit.persentage_two)))
            self.tableWidget.setItem(i, 3, QtWidgets.QTableWidgetItem(str(d.staff_unit.salary)))
            self.tableWidget.setItem(i, 4, QtWidgets.QTableWidgetItem(str(d.division.id)))
            self.tableWidget.setItem(i, 5, QtWidgets.QTableWidgetItem(d.division.name))
            self.tableWidget.setItem(i, 6, QtWidgets.QTableWidgetItem(str(d.division.persentage_one)))
            i += 1

    def show_search_record_dst(self):
        s_roll_no = self.entID.text()
        print(s_roll_no)

        self.tableWidget.clearContents()
        division = self.dst_ctrl.get_division_by_unit_id(s_roll_no)
        self.tableWidget.setRowCount(1)
        self.tableWidget.setItem(0, 0, QtWidgets.QTableWidgetItem(str(s_roll_no)))
        self.tableWidget.setItem(0, 1, QtWidgets.QTableWidgetItem())
        self.tableWidget.setItem(0, 2, QtWidgets.QTableWidgetItem())
        self.tableWidget.setItem(0, 3, QtWidgets.QTableWidgetItem())
        self.tableWidget.setItem(0, 4, QtWidgets.QTableWidgetItem(str(division.id)))
        self.tableWidget.setItem(0, 5, QtWidgets.QTableWidgetItem(division.name))
        self.tableWidget.setItem(0, 6, QtWidgets.QTableWidgetItem(str(division.persentage_one)))

    def calc(self):
        res = self.dst_ctrl.calculate_salary_for_unit(get_selected_record(self.tableWidget))
        self.lblSal.setText(res)

    def set_ent_dst(self):
        row = self.tableWidget.currentRow()
        staffId = self.tableWidget.item(row, 0).text()
        divId = self.tableWidget.item(row, 4).text()
        self.entStaffID.setText(staffId)
        self.entDivID.setText(divId)

    ####================================================================DIV
    def xlsx_export_div(self):
        wb = openpyxl.load_workbook('C:\\Users\\karina\\PycharmProjects\\staffing\\reports\\All.xlsx')
        if 'Divisions' not in wb.sheetnames:
            wb.create_sheet('Divisions')
        ws = wb.get_sheet_by_name('Divisions')
        ws.delete_cols(1, 4)
        ws.delete_rows(1, 100)
        i = 1
        ws.cell(row=i, column=1).value = "ID"
        ws.cell(row=i, column=2).value = "Persent1"
        ws.cell(row=i, column=3).value = "Type"
        ws.cell(row=i, column=4).value = "Name"
        divisions = self.div_ctrl.all_divisions()
        for div in divisions:
            ws.cell(row=i + 1, column=1).value = div.id
            ws.cell(row=i + 1, column=2).value = div.persentage_one
            ws.cell(row=i + 1, column=3).value = div.type
            ws.cell(row=i + 1, column=4).value = div.name
            i += 1
        wb.save('C:\\Users\\karina\\PycharmProjects\\staffing\\reports\\All.xlsx')
        print("xlsx successful")

    def docx_export_div(self):
        document = Document()
        document.add_heading("Divisions", 0)
        table = document.add_table(rows=1, cols=4)
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'ID'
        hdr_cells[1].text = 'Name'
        hdr_cells[2].text = 'Persentage per irregular day'
        hdr_cells[3].text = 'Type'
        divisions = self.div_ctrl.all_divisions()
        for div in divisions:
            row_cells = table.add_row().cells
            row_cells[0].text = str(div.id)
            row_cells[1].text = div.name
            row_cells[2].text = str(div.persentage_one)
            row_cells[3].text = div.type
        document.save('C:\\Users\\karina\\PycharmProjects\\staffing\\reports\\Divisions.docx')
        print("docx successful")

    def register_division(self):

        name = self.entName.text()  # Retrieving entered first name
        pers = self.entPer.text()  # Retrieving entered last name
        type = self.entType.text()  # Retrieving entered contact number

        try:
            self.div_ctrl.create(Division(name, pers, type))
            self.load_divisions()
        except Exception as err:
            print(err)

    def set_ent_div(self):
        row = self.tableWidget_3.currentRow()
        pers = self.tableWidget_3.item(row, 1).text()
        type = self.tableWidget_3.item(row, 2).text()
        name = self.tableWidget_3.item(row, 3).text()
        self.entName.setText(name)
        self.entPer.setText(pers)
        self.entType.setText(type)

    def update_division(self):
        name = self.entName.text()
        pers = self.entPer.text()
        type = self.entType.text()
        print(get_selected_record(self.tableWidget_3))
        self.div_ctrl.update(get_selected_record(self.tableWidget_3), Division(name, pers, type))
        self.load_divisions()

    def delete_division(self):
        self.div_ctrl.delete(get_selected_record(self.tableWidget_3))
        self.load_divisions()
        self.clear_form_div()

    def clear_form_div(self):
        self.entName.clear()
        self.entPer.clear()
        self.entType.clear()

    def load_divisions(self):
        self.tableWidget_3.clearContents()
        divisions = self.div_ctrl.all_divisions()
        i = 0
        self.tableWidget_3.setRowCount(divisions.__len__())
        for div in divisions:
            self.tableWidget_3.setItem(i, 0, QtWidgets.QTableWidgetItem(str(div.id)))
            self.tableWidget_3.setItem(i, 1, QtWidgets.QTableWidgetItem(str(div.persentage_one)))
            self.tableWidget_3.setItem(i, 2, QtWidgets.QTableWidgetItem(str(div.type)))
            self.tableWidget_3.setItem(i, 3, QtWidgets.QTableWidgetItem(str(div.name)))
            i += 1

    def show_search_record_div(self):
        s_roll_no = self.entID_3.text()  # Retrieving entered first name
        division = self.div_ctrl.get_by_id(s_roll_no)
        self.tableWidget_3.clearContents()
        self.tableWidget_3.setRowCount(1)
        self.tableWidget_3.setItem(0, 0, QtWidgets.QTableWidgetItem(str(division.id)))
        self.tableWidget_3.setItem(0, 1, QtWidgets.QTableWidgetItem(division.name))
        self.tableWidget_3.setItem(0, 2, QtWidgets.QTableWidgetItem(str(division.persentage_one)))
        self.tableWidget_3.setItem(0, 3, QtWidgets.QTableWidgetItem(division.type))

        ####================================================================STF

    ####================================================================STF
    def xlsx_export_stf(self):
        wb = openpyxl.load_workbook('C:\\Users\\karina\\PycharmProjects\\staffing\\reports\\All.xlsx')
        if 'Staff units' not in wb.sheetnames:
            wb.create_sheet('Staff units')
        ws = wb.get_sheet_by_name('Staff units')
        ws.delete_cols(1, 5)
        ws.delete_rows(1, 100)
        i = 1
        ws.cell(row=i, column=1).value = "ID"
        ws.cell(row=i, column=2).value = "Persent2"
        ws.cell(row=i, column=3).value = "Vacation"
        ws.cell(row=i, column=4).value = "Position"
        ws.cell(row=i, column=5).value = "Salary"
        units = self.stf_ctrl.all_units()
        for u in units:
            ws.cell(row=i + 1, column=1).value = u.id
            ws.cell(row=i + 1, column=2).value = u.persentage_two
            ws.cell(row=i + 1, column=3).value = u.vacation
            ws.cell(row=i + 1, column=4).value = u.position
            ws.cell(row=i + 1, column=5).value = u.salary
            i += 1
        wb.save('C:\\Users\\karina\\PycharmProjects\\staffing\\reports\\All.xlsx')
        print("xlsx successful")

    def docx_export_stf(self):
        document = Document()
        document.add_heading("Staff units", 0)
        table = document.add_table(rows=1, cols=5)
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'ID'
        hdr_cells[1].text = 'Persent for harmful conditions'
        hdr_cells[2].text = 'Vacation'
        hdr_cells[3].text = 'Position'
        hdr_cells[4].text = 'Salary'
        units = self.stf_ctrl.all_units()
        for u in units:
            row_cells = table.add_row().cells
            row_cells[0].text = str(u.id)
            row_cells[1].text = str(u.persentage_two)
            row_cells[2].text = str(u.vacation)
            row_cells[3].text = u.position
            row_cells[4].text = str(u.salary)
        document.save('C:\\Users\\karina\\PycharmProjects\\staffing\\reports\\Staff Units.docx')
        print("docx successful")

    def register_staff_unit(self):

        position = self.entPos.text()
        pers = self.entPers1.text()
        sal = self.entSal.text()
        vac = self.entVac.text()

        try:
            self.stf_ctrl.create(StaffUnits(pers, vac, position, sal))
            self.load_staff_units()
        except Exception as err:
            print(err)

    def update_staff_unit(self):
        position = self.entPos.text()
        pers = self.entPers1.text()
        sal = self.entSal.text()
        vac = self.entVac.text()
        print(get_selected_record(self.tableWidget_8))
        self.stf_ctrl.update(get_selected_record(self.tableWidget_8), StaffUnits(pers, vac, position, sal))
        self.load_staff_units()

    def delete_staff_unit(self):

        self.stf_ctrl.delete(get_selected_record(self.tableWidget_8))
        self.load_staff_units()
        self.clear_form_stf()

    def clear_form_stf(self):
        self.entPos.clear()
        self.entPers1.clear()
        self.entSal.clear()
        self.entVac.clear()

    def load_staff_units(self):

        self.tableWidget_8.clearContents()
        units = self.stf_ctrl.all_units()
        i = 0
        self.tableWidget_8.setRowCount(units.__len__())
        for u in units:
            self.tableWidget_8.setItem(i, 0, QtWidgets.QTableWidgetItem(str(u.id)))
            self.tableWidget_8.setItem(i, 1, QtWidgets.QTableWidgetItem(str(u.position)))
            self.tableWidget_8.setItem(i, 2, QtWidgets.QTableWidgetItem(str(u.persentage_two)))
            self.tableWidget_8.setItem(i, 3, QtWidgets.QTableWidgetItem(str(u.salary)))
            self.tableWidget_8.setItem(i, 4, QtWidgets.QTableWidgetItem(str(u.vacation)))
            i += 1

    def show_search_record_stf(self):
        s_roll_no = self.entID_8.text()
        stf = self.stf_ctrl.get_by_id(s_roll_no)
        self.tableWidget_8.clearContents()
        self.tableWidget_8.setRowCount(1)
        self.tableWidget_8.setItem(0, 0, QtWidgets.QTableWidgetItem(str(stf.id)))
        self.tableWidget_8.setItem(0, 1, QtWidgets.QTableWidgetItem(stf.position))
        self.tableWidget_8.setItem(0, 2, QtWidgets.QTableWidgetItem(str(stf.persentage_two)))
        self.tableWidget_8.setItem(0, 3, QtWidgets.QTableWidgetItem(str(stf.salary)))
        self.tableWidget_8.setItem(0, 4, QtWidgets.QTableWidgetItem(str(stf.vacation)))

    def set_ent_stf(self):
        row = self.tableWidget_8.currentRow()
        position = self.tableWidget_8.item(row, 1).text()
        pers = self.tableWidget_8.item(row, 2).text()
        sal = self.tableWidget_8.item(row, 3).text()
        vac = self.tableWidget_8.item(row, 4).text()

        self.entPos.setText(position)
        self.entPers1.setText(pers)
        self.entSal.setText(sal)
        self.entVac.setText(vac)


if __name__ == '__main__':
    app = QtWidgets.QApplication(sys.argv)
    window = MainWindow()
    window.show()
    app.exec_()
